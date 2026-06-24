import logging
# pyrefly: ignore [missing-import]
import docx

logger = logging.getLogger(__name__)


def extract_text_from_docx(file_path: str) -> str:
    """Extract raw plain text from a DOCX file using python-docx.

    Iterates over all paragraphs and joins them with newlines, and iterates
    over all tables, rows, and cells to format table content with tabs and newlines.

    Args:
        file_path: Absolute or relative path to the DOCX file.

    Returns:
        The extracted text as a single string, or an empty string on failure.
        Never raises exceptions.
    """
    logger.info("Starting text extraction from DOCX: '%s'", file_path)
    try:
        document = docx.Document(file_path)

        # Paragraphs
        paragraph_text = "\n".join(p.text for p in document.paragraphs)

        # Tables
        table_texts = []
        for table in document.tables:
            row_texts = []
            for row in table.rows:
                row_texts.append("\t".join(cell.text for cell in row.cells))
            table_texts.append("\n".join(row_texts))

        combined_table_text = "\n\n".join(table_texts)

        # Append table text after paragraph text with a double newline separator
        if paragraph_text and combined_table_text:
            combined_text = paragraph_text + "\n\n" + combined_table_text
        elif combined_table_text:
            combined_text = combined_table_text
        else:
            combined_text = paragraph_text

        return combined_text

    except Exception as exc:
        logger.error(
            "Error extracting text from DOCX '%s': %s",
            file_path,
            exc,
        )
        return ""
